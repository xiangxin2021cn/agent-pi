# /// script
# requires-python = ">=3.12"
# dependencies = ["markitdown>=0.1.5,<0.2", "python-docx>=1.1,<2", "diff-match-patch>=20241021", "click>=8.3,<9"]
# ///
"""Document comparison tool.

Converts two documents to Markdown using markitdown, then computes and
displays differences using diff-match-patch.

Usage:
    uv run doc_diff.py [OPTIONS] FILE1 FILE2
"""

import difflib
import sys
import warnings
from pathlib import Path

# Suppress pydub/ffmpeg warning from markitdown[all] — irrelevant for document conversion
warnings.filterwarnings("ignore", message="Couldn't find ffmpeg", category=RuntimeWarning)

import click
from diff_match_patch import diff_match_patch


def write_output(text: str, output_path: str | None) -> None:
    """Write text to file or stdout."""
    if output_path:
        Path(output_path).write_text(text, encoding="utf-8")
        click.echo(f"Output written to {output_path}", err=True)
    else:
        click.echo(text)


def _extract_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def convert_to_text(file_path: str) -> str:
    """Convert a file to text/markdown. Plain text files are read directly."""
    path = Path(file_path)
    ext = path.suffix.lower()

    # For plain text, read directly
    plain_text_exts = {".txt", ".md", ".rst", ".csv", ".tsv", ".log"}
    if ext in plain_text_exts:
        return path.read_text(encoding="utf-8", errors="replace")

    # DOCX fallback path independent of MarkItDown native dependencies.
    if ext == ".docx":
        return _extract_docx_text(path)

    try:
        from markitdown import MarkItDown  # Lazy import to keep CLI startup resilient
    except Exception as e:
        raise click.ClickException(
            "MarkItDown backend unavailable for non-text files. "
            "Install Microsoft Visual C++ Redistributable and retry. "
            f"Details: {e}"
        ) from e

    converter = MarkItDown()
    result = converter.convert(str(path))
    return result.text_content or ""


def format_unified(text1: str, text2: str, name1: str, name2: str) -> str:
    """Generate a unified diff."""
    lines1 = text1.splitlines()
    lines2 = text2.splitlines()
    diff = difflib.unified_diff(lines1, lines2, fromfile=name1, tofile=name2, lineterm="")
    return "\n".join(diff)


def format_side_by_side(text1: str, text2: str, name1: str, name2: str, width: int = 80) -> str:
    """Generate a side-by-side diff."""
    lines1 = text1.splitlines()
    lines2 = text2.splitlines()

    col_width = (width - 3) // 2  # 3 for ' | ' separator
    header = f"{name1:<{col_width}} | {name2}"
    separator = "-" * col_width + "-+-" + "-" * col_width

    result_lines = [header, separator]

    # Use difflib to align the lines
    matcher = difflib.SequenceMatcher(None, lines1, lines2)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for i in range(i1, i2):
                left = lines1[i][:col_width].ljust(col_width)
                right = lines2[j1 + (i - i1)][:col_width]
                result_lines.append(f"{left} | {right}")
        elif tag == "replace":
            max_len = max(i2 - i1, j2 - j1)
            for k in range(max_len):
                left = lines1[i1 + k][:col_width].ljust(col_width) if (i1 + k) < i2 else " " * col_width
                right = lines2[j1 + k][:col_width] if (j1 + k) < j2 else ""
                marker = "*"
                result_lines.append(f"{left} {marker} {right}")
        elif tag == "delete":
            for i in range(i1, i2):
                left = lines1[i][:col_width].ljust(col_width)
                result_lines.append(f"{left} < ")
        elif tag == "insert":
            for j in range(j1, j2):
                left = " " * col_width
                right = lines2[j][:col_width]
                result_lines.append(f"{left} > {right}")

    return "\n".join(result_lines)


def format_summary(text1: str, text2: str, name1: str, name2: str, word_level: bool = False) -> str:
    """Generate a summary of differences."""
    dmp = diff_match_patch()

    diffs = dmp.diff_main(text1, text2)
    dmp.diff_cleanupSemantic(diffs)

    lines1 = text1.splitlines()
    lines2 = text2.splitlines()

    # Count changes
    insertions = 0
    deletions = 0
    insert_chars = 0
    delete_chars = 0

    for op, data in diffs:
        if op == 1:  # Insert
            insertions += 1
            insert_chars += len(data)
        elif op == -1:  # Delete
            deletions += 1
            delete_chars += len(data)

    # Line-level stats
    line_matcher = difflib.SequenceMatcher(None, lines1, lines2)
    added_lines = 0
    removed_lines = 0
    changed_lines = 0
    for tag, i1, i2, j1, j2 in line_matcher.get_opcodes():
        if tag == "insert":
            added_lines += j2 - j1
        elif tag == "delete":
            removed_lines += i2 - i1
        elif tag == "replace":
            changed_lines += max(i2 - i1, j2 - j1)

    similarity = line_matcher.ratio() * 100

    summary_parts = [
        f"Comparison: {name1} vs {name2}",
        f"",
        f"File 1: {len(lines1)} lines, {len(text1)} characters",
        f"File 2: {len(lines2)} lines, {len(text2)} characters",
        f"",
        f"Similarity: {similarity:.1f}%",
        f"",
        f"Line changes:",
        f"  Added lines:   {added_lines}",
        f"  Removed lines: {removed_lines}",
        f"  Changed lines: {changed_lines}",
        f"",
        f"Character-level changes:",
        f"  Insertions: {insertions} ({insert_chars} characters)",
        f"  Deletions:  {deletions} ({delete_chars} characters)",
    ]

    if word_level:
        # Show word-level changes with markers
        summary_parts.append("")
        summary_parts.append("Word-level changes ([-deleted-] {+inserted+}):")
        summary_parts.append("")
        change_parts: list[str] = []
        for op, data in diffs:
            if op == 0:
                change_parts.append(data)
            elif op == -1:
                change_parts.append(f"[-{data}-]")
            elif op == 1:
                change_parts.append(f"{{+{data}+}}")
        summary_parts.append("".join(change_parts))

    return "\n".join(summary_parts)


@click.command()
@click.argument("file1", type=click.Path(exists=True, dir_okay=False))
@click.argument("file2", type=click.Path(exists=True, dir_okay=False))
@click.option(
    "--format",
    "fmt",
    type=click.Choice(["unified", "side-by-side", "summary"]),
    default="unified",
    help="Diff output format (default: unified).",
)
@click.option("--word-level", is_flag=True, default=False, help="Show word-level differences (summary format).")
@click.option("-o", "--output", type=click.Path(), default=None, help="Write output to file.")
def main(file1: str, file2: str, fmt: str, word_level: bool, output: str | None) -> None:
    """Compare two documents by converting them to Markdown first.

    Supports any format that markitdown can handle: .docx, .xlsx, .pptx, .pdf,
    .html, .ipynb, .xml, .txt, .md, and more.
    """
    name1 = Path(file1).name
    name2 = Path(file2).name

    try:
        click.echo(f"Converting {name1}...", err=True)
        text1 = convert_to_text(file1)

        click.echo(f"Converting {name2}...", err=True)
        text2 = convert_to_text(file2)
    except Exception as e:
        click.echo(f"Error converting files: {e}", err=True)
        sys.exit(1)

    if text1 == text2:
        write_output("Files are identical.", output)
        return

    if fmt == "unified":
        result = format_unified(text1, text2, name1, name2)
    elif fmt == "side-by-side":
        result = format_side_by_side(text1, text2, name1, name2)
    elif fmt == "summary":
        result = format_summary(text1, text2, name1, name2, word_level=word_level)
    else:
        result = format_unified(text1, text2, name1, name2)

    if not result.strip():
        write_output("No differences found (content is identical after conversion).", output)
    else:
        write_output(result, output)


if __name__ == "__main__":
    main()
