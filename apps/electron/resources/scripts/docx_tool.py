# /// script
# requires-python = ">=3.12"
# dependencies = ["python-docx>=1.2,<2", "click>=8.3,<9"]
# ///
"""Word document (.docx) creation and editing tool.

Commands: create, template, info, replace, extract.

Usage:
    uv run docx_tool.py COMMAND [OPTIONS]
"""

import json
import re
import sys
from pathlib import Path

import click
from docx import Document
from docx.shared import Pt


def write_output(text: str, output_path: str | None) -> None:
    """Write text to file or stdout."""
    if output_path:
        Path(output_path).write_text(text, encoding="utf-8")
        click.echo(f"Output written to {output_path}", err=True)
    else:
        click.echo(text)


def markdown_to_docx(md_text: str, doc: Document) -> None:
    """Convert simple markdown text to docx paragraphs.

    Supports: # headings (h1-h6), **bold**, *italic*, - bullet lists,
    1. numbered lists, blank lines as paragraph breaks, and --- as page breaks.
    """
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Page break
        if stripped in ("---", "***", "___"):
            doc.add_page_break()
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet_match:
            text = bullet_match.group(1)
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline_formatting(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if num_match:
            text = num_match.group(1)
            p = doc.add_paragraph(style="List Number")
            _apply_inline_formatting(p, text)
            i += 1
            continue

        # Empty line - skip
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _apply_inline_formatting(p, stripped)
        i += 1


def _apply_inline_formatting(paragraph, text: str) -> None:
    """Apply bold and italic inline formatting to a paragraph."""
    # Clear any default runs
    paragraph.clear()

    # Pattern to match **bold**, *italic*, ***bold italic***
    pattern = r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+|\*(?!\*)))"

    for match in re.finditer(pattern, text):
        if match.group(2):  # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # **bold**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # plain text
            paragraph.add_run(match.group(5))


@click.group()
def cli() -> None:
    """Word document (.docx) creation and editing tool."""
    pass


@cli.command()
@click.option("--from-file", type=click.Path(exists=True, dir_okay=False), default=None, help="Input text or markdown file.")
@click.option("--text", type=str, default=None, help="Direct text content (supports basic markdown).")
@click.option("--title", type=str, default=None, help="Document title.")
@click.option("--font-size", type=float, default=11, help="Base font size in points (default: 11).")
@click.option("-o", "--output", type=click.Path(), required=True, help="Output .docx file path.")
def create(from_file: str | None, text: str | None, title: str | None, font_size: float, output: str) -> None:
    """Create a new Word document from text or markdown.

    Provide content via --from-file (reads a .txt/.md file) or --text (inline string).
    Basic markdown formatting is supported: headings, bold, italic, lists.
    """
    if from_file is None and text is None:
        click.echo("Error: provide --from-file or --text.", err=True)
        sys.exit(1)

    try:
        doc = Document()

        # Set default font size
        style = doc.styles["Normal"]
        style.font.size = Pt(font_size)

        if title:
            doc.add_heading(title, level=0)

        if from_file:
            content = Path(from_file).read_text(encoding="utf-8")
        else:
            content = text or ""

        markdown_to_docx(content, doc)

        doc.save(output)
        click.echo(f"Document created: {output}", err=True)
    except Exception as e:
        click.echo(f"Error: {e}", err=True)
        sys.exit(1)


@cli.command()
@click.argument("template_file", type=click.Path(exists=True, dir_okay=False))
@click.option("--data", type=str, required=True, help="JSON string or path to JSON file with template values.")
@click.option("-o", "--output", type=click.Path(), required=True, help="Output .docx file path.")
def template(template_file: str, data: str, output: str) -> None:
    """Fill a Word document template with JSON data.

    Replaces {{placeholder}} patterns in the document with values from JSON.
    The JSON should map placeholder names to values, e.g. {"name": "John", "date": "2024-01-01"}.
    """
    try:
        # Parse data
        data_path = Path(data)
        if data_path.exists() and data_path.is_file():
            template_data = json.loads(data_path.read_text(encoding="utf-8"))
        else:
            template_data = json.loads(data)

        if not isinstance(template_data, dict):
            click.echo("Error: JSON data must be an object.", err=True)
            sys.exit(1)

        doc = Document(template_file)

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            _replace_in_paragraph(paragraph, template_data)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, template_data)

        # Replace in headers/footers
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                if not header_footer.is_linked_to_previous:
                    for paragraph in header_footer.paragraphs:
                        _replace_in_paragraph(paragraph, template_data)

        doc.save(output)
        click.echo(f"Template filled and saved to {output}", err=True)
    except json.JSONDecodeError as e:
        click.echo(f"Error parsing JSON: {e}", err=True)
        sys.exit(1)
    except Exception as e:
        click.echo(f"Error: {e}", err=True)
        sys.exit(1)


def _replace_in_paragraph(paragraph, data: dict[str, str]) -> None:
    """Replace {{key}} placeholders in a paragraph while preserving formatting."""
    full_text = paragraph.text
    if "{{" not in full_text:
        return

    # Single-pass replacement to avoid cascading (a replacement value containing
    # {{other_key}} should not be substituted again).
    def _replacer(match: re.Match) -> str:
        key = match.group(1)
        if key in data:
            return str(data[key])
        return match.group(0)  # leave unmatched placeholders as-is

    full_text = re.sub(r"\{\{(\w+)\}\}", _replacer, full_text)

    # Rebuild runs with the replaced text
    if paragraph.runs:
        # Clear all runs except the first, put full text in first run
        for i in range(len(paragraph.runs) - 1, 0, -1):
            paragraph.runs[i].text = ""
        paragraph.runs[0].text = full_text


@cli.command()
@click.argument("file", type=click.Path(exists=True, dir_okay=False))
@click.option("-o", "--output", type=click.Path(), default=None, help="Write output to file.")
def info(file: str, output: str | None) -> None:
    """Show document information: paragraphs, sections, tables, styles."""
    try:
        doc = Document(file)

        # Count elements
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        section_count = len(doc.sections)

        # Styles used
        styles_used = set()
        for p in doc.paragraphs:
            if p.style and p.style.name:
                styles_used.add(p.style.name)

        # Word count
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)

        # Core properties
        props = doc.core_properties
        info_dict: dict[str, object] = {
            "file": str(Path(file).resolve()),
            "paragraphs": paragraph_count,
            "tables": table_count,
            "sections": section_count,
            "word_count": word_count,
            "styles_used": sorted(styles_used),
            "properties": {
                "title": props.title,
                "author": props.author,
                "subject": props.subject,
                "created": str(props.created) if props.created else None,
                "modified": str(props.modified) if props.modified else None,
                "last_modified_by": props.last_modified_by,
                "revision": props.revision,
            },
        }

        # Template placeholders
        placeholders = set()
        for p in doc.paragraphs:
            for match in re.finditer(r"\{\{(\w+)\}\}", p.text):
                placeholders.add(match.group(1))
        if placeholders:
            info_dict["template_placeholders"] = sorted(placeholders)

        write_output(json.dumps(info_dict, indent=2, default=str), output)
    except Exception as e:
        click.echo(f"Error: {e}", err=True)
        sys.exit(1)


@cli.command()
@click.argument("file", type=click.Path(exists=True, dir_okay=False))
@click.option("--find", type=str, required=True, help="Text to find.")
@click.option("--replace-with", type=str, required=True, help="Replacement text.")
@click.option("--case-sensitive/--no-case-sensitive", default=True, help="Case-sensitive matching (default: yes).")
@click.option("-o", "--output", type=click.Path(), required=True, help="Output .docx file path.")
def replace(file: str, find: str, replace_with: str, case_sensitive: bool, output: str) -> None:
    """Find and replace text in a Word document."""
    try:
        doc = Document(file)
        count = 0

        for paragraph in doc.paragraphs:
            c = _find_replace_paragraph(paragraph, find, replace_with, case_sensitive)
            count += c

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        c = _find_replace_paragraph(paragraph, find, replace_with, case_sensitive)
                        count += c

        # Replace in headers/footers
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                if not header_footer.is_linked_to_previous:
                    for paragraph in header_footer.paragraphs:
                        c = _find_replace_paragraph(paragraph, find, replace_with, case_sensitive)
                        count += c

        doc.save(output)
        click.echo(f"Replaced {count} occurrence(s). Saved to {output}", err=True)
    except Exception as e:
        click.echo(f"Error: {e}", err=True)
        sys.exit(1)


def _find_replace_paragraph(paragraph, find: str, replace_with: str, case_sensitive: bool) -> int:
    """Replace text in a paragraph, returns count of replacements."""
    full_text = paragraph.text
    check_text = full_text if case_sensitive else full_text.lower()
    find_text = find if case_sensitive else find.lower()

    if find_text not in check_text:
        return 0

    count = check_text.count(find_text)

    if case_sensitive:
        new_text = full_text.replace(find, replace_with)
    else:
        # Case-insensitive replace
        pattern = re.compile(re.escape(find), re.IGNORECASE)
        new_text = pattern.sub(replace_with, full_text)

    # Rebuild runs
    if paragraph.runs:
        for i in range(len(paragraph.runs) - 1, 0, -1):
            paragraph.runs[i].text = ""
        paragraph.runs[0].text = new_text

    return count


@cli.command()
@click.argument("file", type=click.Path(exists=True, dir_okay=False))
@click.option("--include-tables/--no-tables", default=True, help="Include table content (default: yes).")
@click.option("-o", "--output", type=click.Path(), default=None, help="Write output to file.")
def extract(file: str, include_tables: bool, output: str | None) -> None:
    """Extract all text content from a Word document."""
    try:
        doc = Document(file)
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Add heading markers
                if paragraph.style and paragraph.style.name.startswith("Heading"):
                    try:
                        level = int(paragraph.style.name.split()[-1])
                        text = "#" * level + " " + text
                    except (ValueError, IndexError):
                        pass
                parts.append(text)

        if include_tables:
            for t_idx, table in enumerate(doc.tables):
                parts.append(f"\n[Table {t_idx + 1}]")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append(" | ".join(cells))

        write_output("\n".join(parts), output)
    except Exception as e:
        click.echo(f"Error: {e}", err=True)
        sys.exit(1)


if __name__ == "__main__":
    cli()
