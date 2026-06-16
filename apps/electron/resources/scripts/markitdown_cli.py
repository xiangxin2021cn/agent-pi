# /// script
# requires-python = ">=3.12"
# dependencies = ["markitdown[pdf,pptx,docx,xlsx]>=0.1.5,<0.2", "python-docx>=1.1,<2", "click>=8.3,<9"]
# ///
"""Universal document to Markdown converter.

Converts .docx, .xlsx, .pptx, .pdf, .html, .ipynb, .xml, .rss, .zip, .msg
and other supported formats to Markdown using the markitdown library.

Usage:
    uv run markitdown_cli.py [OPTIONS] FILE
"""

import sys
import warnings
from pathlib import Path

# Suppress pydub/ffmpeg warning from markitdown[all] — irrelevant for document conversion
warnings.filterwarnings("ignore", message="Couldn't find ffmpeg", category=RuntimeWarning)

import click


SUPPORTED_EXTENSIONS = {
    ".docx",
    ".xlsx",
    ".pptx",
    ".pdf",
    ".html",
    ".htm",
    ".ipynb",
    ".xml",
    ".rss",
    ".atom",
    ".zip",
    ".msg",
    ".eml",
    ".csv",
    ".tsv",
    ".json",
    ".txt",
    ".md",
    ".rst",
    ".rtf",
    ".jpg",
    ".jpeg",
    ".png",
    ".gif",
    ".bmp",
    ".tiff",
    ".wav",
    ".mp3",
}


PLAIN_TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".csv", ".tsv", ".json", ".xml", ".html", ".htm", ".log"}


def _extract_docx_text(file_path: Path) -> str:
    """Fallback DOCX extraction that avoids MarkItDown native dependency chain."""
    from docx import Document

    doc = Document(str(file_path))
    parts: list[str] = []

    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


def write_output(text: str, output_path: str | None) -> None:
    """Write text to file or stdout."""
    if output_path:
        Path(output_path).write_text(text, encoding="utf-8")
        click.echo(f"Output written to {output_path}", err=True)
    else:
        click.echo(text)


@click.command()
@click.argument("file", type=click.Path(exists=True, dir_okay=False))
@click.option("-o", "--output", type=click.Path(), default=None, help="Write output to file instead of stdout.")
def main(file: str, output: str | None) -> None:
    """Convert a document to Markdown.

    Supports .docx, .xlsx, .pptx, .pdf, .html, .ipynb, .xml, .rss, .zip, .msg,
    and many other formats.
    """
    file_path = Path(file)
    ext = file_path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        click.echo(
            f"Warning: extension '{ext}' may not be supported. Attempting conversion anyway.",
            err=True,
        )

    # Fast fallback for plain text-like inputs.
    if ext in PLAIN_TEXT_EXTENSIONS:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        write_output(text, output)
        return

    # DOCX fallback path that avoids MarkItDown's optional native stack.
    if ext == ".docx":
        try:
            write_output(_extract_docx_text(file_path), output)
            return
        except Exception:
            # Fall back to MarkItDown path below if python-docx extraction fails unexpectedly.
            pass

    try:
        from markitdown import MarkItDown  # Lazy import: avoids startup crashes on systems missing optional native deps

        converter = MarkItDown()
        result = converter.convert(str(file_path))
        write_output(result.text_content or "", output)
    except Exception as e:
        click.echo(
            f"Error converting {file_path.name}: {e}\n"
            "If this is a native dependency issue (e.g. onnxruntime DLL), "
            "install Microsoft Visual C++ Redistributable and retry.",
            err=True,
        )
        sys.exit(1)


if __name__ == "__main__":
    main()
